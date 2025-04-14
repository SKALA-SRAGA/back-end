import json
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from .exchange import get_exchange_rate_safe

HEADER_COLOR = "4472C4"
SUBHEADER_COLOR = "D9E1F2"
BORDER_COLOR = "BFBFBF"


def create_expense_report(json_file_path, output_path):

    print(f"JSON 데이터 파일: {json_file_path}")
    print(f"출력 파일 경로: {output_path}")

    with open(json_file_path, "r", encoding="utf-8") as f:
        json_data = json.load(f)

    expenses = []
    if isinstance(json_data, list):
        for file_entry in json_data:
            file_name = file_entry.get("file", "")
            results = file_entry.get("result", [])
            for result in results:
                expense_item = result.copy()
                expense_item["file"] = file_name
                expenses.append(expense_item)

    categorized_expenses = categorize_expenses(expenses)
    doc = Document()
    set_document_style(doc)
    add_title(doc, "출장경비 내역서")
    add_basic_info_table(doc)
    add_expense_table(doc, categorized_expenses)
    doc.save(output_path)
    print(f"출장경비 내역서가 생성되었습니다: {output_path}")


def set_document_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)

    title_style = doc.styles["Title"]
    title_style.font.name = "맑은 고딕"
    title_style.font.size = Pt(18)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_title(doc, title_text):
    title = doc.add_heading(title_text, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()


def add_basic_info_table(doc):
    heading = doc.add_heading("출장내역", level=2)
    heading.style.font.size = Pt(14)
    heading.style.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    basic_table = doc.add_table(rows=6, cols=2)
    basic_table.style = "Table Grid"
    basic_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(basic_table, 16.0)
    set_column_width(basic_table, 0, 2.0)
    set_column_width(basic_table, 1, 6.0)
    set_table_borders(basic_table)

    info_labels = ["출장자", "부서", "기간", "목적", "장소", "동반자"]
    for i, label in enumerate(info_labels):
        row = basic_table.rows[i]
        cell = row.cells[0]
        cell.text = label
        apply_cell_shading(cell, SUBHEADER_COLOR)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_vertical_alignment(cell, WD_CELL_VERTICAL_ALIGNMENT.CENTER)

    doc.add_paragraph()


def format_exchange_rate(currency: str, rate: int) -> str:
    if currency.upper() == "JPY":
        return f"{rate * 100}"
    else:
        return f"{rate}"


def add_expense_table(doc, categorized_expenses):
    heading = doc.add_heading("경비내역", level=2)
    heading.style.font.size = Pt(14)
    heading.style.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    total_items = sum(len(items) for items in categorized_expenses.values())
    table = doc.add_table(rows=total_items + 2, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    set_table_width(table, 16.0)
    col_widths = [2.0, 4.0, 3.0, 2.5, 2.0, 2.5]
    for i, width in enumerate(col_widths):
        set_column_width(table, i, width)
    set_table_borders(table)

    header_cells = table.rows[0].cells
    headers = ["항목", "내용", "출장국가\n화폐금액", "원화\n금액", "환율", "날짜"]
    for i, header in enumerate(headers):
        cell = header_cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(header)
        run.font.size = Pt(10)
        apply_cell_shading(cell, SUBHEADER_COLOR)
        set_cell_vertical_alignment(cell, WD_CELL_VERTICAL_ALIGNMENT.CENTER)

    category_order = ["식비", "교통비", "숙박비", "기타"]
    row_idx = 1
    for category in category_order:
        expenses = categorized_expenses.get(category, [])
        for expense in expenses:
            row = table.rows[row_idx]

            currency = expense.get("currency", "KRW")
            amount = expense.get("amount", 0)
            exchange_rate = get_exchange_rate_safe(currency)
            krw_amount = calculate_krw_amount(amount, currency, exchange_rate)

            row.cells[0].text = category
            row.cells[1].text = (
                expense.get("vendor_kor", "") + " - " + expense.get("details", "")
            )
            row.cells[2].text = f"{currency} {int(amount):,}"
            row.cells[3].text = f"{int(krw_amount):,}"
            row.cells[4].text = format_exchange_rate(currency, exchange_rate)
            row.cells[5].text = expense.get("date", "")

            for cell in row.cells:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(cell, WD_CELL_VERTICAL_ALIGNMENT.CENTER)

            row_idx += 1

    total_row = table.rows[-1]
    total_row.cells[0].text = "합계"
    total_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_row.cells[1].text = ""

    total_krw = 0
    for expenses in categorized_expenses.values():
        for expense in expenses:
            currency = expense.get("currency", "KRW")
            amount = expense.get("amount", 0)
            exchange_rate = get_exchange_rate_safe(currency)
            total_krw += calculate_krw_amount(amount, currency, exchange_rate)

    total_row.cells[3].text = f"{int(total_krw):,}"
    total_row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_vertical_alignment(total_row.cells[3], WD_CELL_VERTICAL_ALIGNMENT.CENTER)


def calculate_krw_amount(amount, currency, exchange_rate=None):
    if currency == "KRW":
        return int(amount)
    if exchange_rate is None:
        exchange_rate = get_exchange_rate_safe(currency)
    return int(float(amount) * exchange_rate)


def categorize_expenses(expenses):
    categories = {"식비": [], "교통비": [], "숙박비": [], "기타": []}
    for expense in expenses:
        category = expense.get("category", "기타").lower().strip()
        if category in categories:
            categories[category].append(expense)
        else:
            categories["기타"].append(expense)
    return categories


def set_table_width(table, width_in_cm):
    table.autofit = False
    table._tbl.xpath("./w:tblPr/w:tblW")[0].set(qn("w:w"), str(width_in_cm * 567))
    table._tbl.xpath("./w:tblPr/w:tblW")[0].set(qn("w:type"), "dxa")


def set_column_width(table, column_index, width_in_cm):
    grid_col = table._tbl.xpath("./w:tblGrid/w:gridCol")[column_index]
    grid_col.set(qn("w:w"), str(int(width_in_cm * 567)))
    for row in table.rows:
        try:
            cell = row.cells[column_index]
            cell_width = OxmlElement("w:tcW")
            cell_width.set(qn("w:w"), str(int(width_in_cm * 567)))
            cell_width.set(qn("w:type"), "dxa")
            cell._tc.get_or_add_tcPr().append(cell_width)
        except IndexError:
            pass


def set_table_borders(table, border_size=1, border_color=BORDER_COLOR):
    tbl = table._tbl
    tblPr = tbl.xpath("./w:tblPr")[0]
    tblBorders = OxmlElement("w:tblBorders")
    for border_pos in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_pos}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(border_size * 4))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), border_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def apply_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_vertical_alignment(cell, alignment):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(
        qn("w:val"),
        {
            WD_CELL_VERTICAL_ALIGNMENT.TOP: "top",
            WD_CELL_VERTICAL_ALIGNMENT.CENTER: "center",
            WD_CELL_VERTICAL_ALIGNMENT.BOTTOM: "bottom",
        }[alignment],
    )
    tcPr.append(vAlign)
